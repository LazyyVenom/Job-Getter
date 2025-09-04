from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Create a new document
doc = Document()

# Set default font style
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# ---------- HEADER ----------
p = doc.add_paragraph()
run = p.add_run("Anubhav Choubey")
run.font.size = Pt(20)
run.bold = True
p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

p = doc.add_paragraph("+91 9131XXXXXX | choubey.anubhav253@gmail.com | anubhavchoubey.com")
p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

p = doc.add_paragraph("GitHub | LinkedIn | Kaggle | Jabalpur, Madhya Pradesh")
p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

doc.add_paragraph("_" * 100)

# ---------- SECTION HELPER ----------
def add_section(title):
    p = doc.add_paragraph()
    run = p.add_run(title.upper())
    run.font.size = Pt(13)
    run.bold = True

# ---------- EXPERIENCE ----------
add_section("Experience")

experiences = [
    ("Technical Intern @Rozana.in (Onsite)", "Jul 2025 - Present", [
        "Built a pipeline for automated image generation and editing using SOTA open-source models (Flux Krea Dev, Qwen Image Edit, Vision LLMs).",
        "Automated generation of display names, descriptions, and visual tags using Gemini API, and optimized Typesense queries for improved search relevance.",
        "Developed a pipeline to fetch, process, and deploy geofencing data for store coverage areas, enabling seamless production integration."
    ]),
    ("Deep Learning Intern @Devnex Technologies", "Sep 2024 - Jul 2025", [
        "Deployed open-source models (Flux, Flux Kontext, InternVL, Gemma, etc.) on HuggingFace Endpoints and Modal.com with quantized versions for efficient remote execution.",
        "Optimized model performance by reducing memory usage and runtime through parallel processing, multithreading, and pipelining.",
        "Collaborated in professional environments using ticket-based systems, Git workflows, and maintained thorough code documentation."
    ]),
    ("Intern @Lexi.ai", "May 2024 - Jun 2025", [
        "Developed a Django backend with a web-hosted SQL database, including MySQL schema design and robust validation for accurate data transmission.",
        "Contributed to website basic frontend development using HTML, CSS, JS (React)."
    ]),
    ("Freelancer", "Sep 2023 - Sep 2024", [
        "Completed 15+ projects spanning Python, FastAPI, Django, Machine Learning, automation, and various Python frameworks."
    ])
]

for role, date, bullets in experiences:
    p = doc.add_paragraph()
    run = p.add_run(f"{role} ")
    run.bold = True
    p.add_run(f"{date}").italic = True
    for b in bullets:
        doc.add_paragraph(b, style="List Bullet")

# ---------- PROJECTS ----------
add_section("Projects")

projects = [
    ("School Sphere - School Management System", "Python, FastAPI, Celery, Redis, JWT", [
        "Engineered a secure backend with JWT-based authentication and multi-role access control, ensuring safe request handling and granular read/write permissions.",
        "Optimized performance using Redis-based class-level caching, Celery for task scheduling, and multithreading to reduce redundant processing and improve scalability.",
        "Followed industry best practices (logging, lazy loading, exception handling) while building with Python, FastAPI, Redis, Celery, and JWT for a production-ready system."
    ]),
    ("Gesture Controlled Virtual Assistant", "OpenCV, Python, MediaPipe, NumPy", [
        "Created a virtual assistant that uses gesture control for hands-free computer interaction.",
        "Incorporated facial and hand filters to enhance the representation of the subject."
    ]),
    ("Chess AI", "Python, PyGame, Minimax Algorithm, Artificial Intelligence, Parallel Programming", [
        "Built a Chess game with an AI opponent (~1000 ELO) using Python, PyGame, and the Minimax algorithm.",
        "Optimized AI performance with parallel processing, multithreading, and best coding practices."
    ])
]

for proj, tech, bullets in projects:
    p = doc.add_paragraph()
    run = p.add_run(proj)
    run.bold = True
    p.add_run(f"   ({tech})").italic = True
    for b in bullets:
        doc.add_paragraph(b, style="List Bullet")

# ---------- SKILLS ----------
add_section("Skills")
doc.add_paragraph("Languages & Frameworks: Python, Django, FastAPI, SQL")
doc.add_paragraph("AI/ML & Deep Learning: Prompting, ComfyUI, Pipelining, HuggingFace, CV, NLP, OpenCV, Inferencing")
doc.add_paragraph("Data Handling: MySQL, SQLAlchemy, Pandas, NumPy, BeautifulSoup4, Firebase, REST APIs, Webhooks")
doc.add_paragraph("Cloud & Deployment: HuggingFace Endpoints, AWS, Linux, Docker")
doc.add_paragraph("Tools & Automation: Selenium, Git, GitHub")

# ---------- EDUCATION ----------
add_section("Education")
doc.add_paragraph("B.Tech (Artificial Intelligence and Data Science): JEC (Est. 1947) -- 8.4*   Jul 2022 -- Apr 2026")
doc.add_paragraph("CBSE (Senior Secondary Education): KVOFK -- 92.2%   Apr 2021 -- Jul 2022")

# ---------- ACCOMPLISHMENTS ----------
add_section("Accomplishments")
accomplishments = [
    "Led 'Team Conquerors!' to SIH’24 Finals.",
    "LeetCode: Solved 200+ problems in Data Structures and Algorithms.",
    "Led the winning team at CodeKumbh Hackathon, showcasing leadership and teamwork skills.",
    "Team leader for a Top 4 project at the HackCrux Hackathon.",
    "Built and manage a coding-focused YouTube channel with 9K+ subscribers and 2.5M+ total views.",
    "Certifications: Python, AI, ML, Technical Support Fundamentals (Coursera)."
]
for acc in accomplishments:
    doc.add_paragraph(acc, style="List Bullet")

# Save document
output_path = "Updated_Resume_Anubhav_Choubey.docx"
doc.save(output_path)
output_path

from docx2pdf import convert

# Convert DOCX to PDF
convert("Updated_Resume_Anubhav_Choubey.docx", "Updated_Resume_Anubhav_Choubey.pdf")